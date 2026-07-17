"""
parser.py
---------
Utilities for extracting raw text from uploaded resume / job description
files. Supports PDF, DOCX and plain TXT files.
"""

from __future__ import annotations

import io

import pdfplumber
from docx import Document
from fastapi import UploadFile

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileTypeError(Exception):
    """Raised when a file with an unsupported extension is uploaded."""


class EmptyDocumentError(Exception):
    """Raised when no extractable text could be found in a document."""


def get_extension(filename: str) -> str:
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF using pdfplumber, page by page."""
    text_chunks: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_chunks.append(page_text)
    return "\n".join(text_chunks).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file, including paragraphs and tables."""
    document = Document(io.BytesIO(file_bytes))

    text_chunks: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_chunks.append(cell.text)

    return "\n".join(text_chunks).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode a plain text file, trying a couple of common encodings."""
    for encoding in ("utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise UnsupportedFileTypeError("Could not decode text file with supported encodings.")


async def extract_text(file: UploadFile) -> str:
    """
    Read an UploadFile, dispatch to the correct extractor based on its
    extension, and return the extracted plain text.

    Raises:
        UnsupportedFileTypeError: if the extension isn't pdf/docx/txt.
        EmptyDocumentError: if no text could be extracted.
    """
    if not file.filename:
        raise UnsupportedFileTypeError("Uploaded file is missing a filename.")

    extension = get_extension(file.filename)
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{extension}'. Supported types: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )

    file_bytes = await file.read()
    if not file_bytes:
        raise EmptyDocumentError(f"'{file.filename}' is empty.")

    if extension == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif extension == ".docx":
        text = extract_text_from_docx(file_bytes)
    else:  # .txt
        text = extract_text_from_txt(file_bytes)

    if not text or not text.strip():
        raise EmptyDocumentError(
            f"No readable text could be extracted from '{file.filename}'. "
            "The file may be a scanned image or corrupted."
        )

    return text
